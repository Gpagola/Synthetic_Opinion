"""Generación del informe del focus group y exportación a PDF/Word/Excel."""

from __future__ import annotations

import io

import markdown as md
from docx import Document
from openpyxl import Workbook
from sqlalchemy.orm import Session

from app.models import FocusGroup, Report
from app.services.llm import get_llm

_SYSTEM = """Eres un analista senior de investigación cualitativa. Redactas
informes de focus group claros, accionables y bien estructurados. Citas
textualmente a los participantes cuando aporta valor. Respondes en JSON válido."""


def _transcript(fg: FocusGroup) -> str:
    bloques = []
    for q in fg.questions:
        tipo = "uno a uno" if q.tipo == "uno_a_uno" else "general"
        bloques.append(f"### Pregunta ({tipo}): {q.texto}")
        for r in sorted(q.responses, key=lambda x: x.orden):
            bloques.append(f"- {r.persona_nombre}: {r.texto}")
    return "\n".join(bloques)


def _build_prompt(fg: FocusGroup) -> str:
    return f"""Analiza el siguiente focus group y redacta un informe en {fg.idioma}.

Nombre: {fg.nombre}
Tema: {fg.tema}
Descripción: {fg.descripcion}

Transcripción:
{_transcript(fg)}

Devuelve un objeto JSON con esta forma:
{{
  "markdown": "Informe completo en formato Markdown con estas secciones: Resumen ejecutivo, Temas clave, Consensos y disensos, Segmentos/perfiles, Citas destacadas, Conclusiones y recomendaciones.",
  "metadatos": {{
    "temas": ["string"],
    "sentimiento_global": "string",
    "citas_destacadas": ["string"]
  }}
}}"""


def generate_report(db: Session, fg: FocusGroup) -> Report:
    llm = get_llm()
    data = llm.complete_json(_SYSTEM, _build_prompt(fg))
    contenido = data.get("markdown", "").strip() or "_No se pudo generar el informe._"
    metadatos = data.get("metadatos", {})

    report = Report(
        focus_group_id=fg.id,
        contenido_markdown=contenido,
        metadatos=metadatos,
    )
    db.add(report)
    db.commit()
    db.refresh(report)
    return report


# ---------- Exportadores ----------


def export_pdf(report: Report, fg: FocusGroup) -> bytes:
    from weasyprint import HTML  # import perezoso (dependencias del sistema)

    body_html = md.markdown(report.contenido_markdown, extensions=["extra", "sane_lists"])
    html = f"""<!doctype html><html><head><meta charset="utf-8">
<style>
  body {{ font-family: Arial, sans-serif; margin: 2.5cm; line-height: 1.5; color: #222; }}
  h1 {{ color: #1a3c6e; }} h2, h3 {{ color: #2a5a9e; }}
  blockquote {{ border-left: 3px solid #ccc; margin: 0; padding-left: 1em; color: #555; }}
</style></head>
<body><h1>{fg.nombre}</h1><p><em>{fg.tema}</em></p>{body_html}</body></html>"""
    return HTML(string=html).write_pdf()


def export_docx(report: Report, fg: FocusGroup) -> bytes:
    doc = Document()
    doc.add_heading(fg.nombre, level=0)
    if fg.tema:
        doc.add_paragraph(fg.tema).italic = True

    for line in report.contenido_markdown.splitlines():
        stripped = line.strip()
        if not stripped:
            continue
        if stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=1)
        elif stripped.startswith(("- ", "* ")):
            doc.add_paragraph(stripped[2:], style="List Bullet")
        else:
            doc.add_paragraph(stripped)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def export_xlsx(fg: FocusGroup) -> bytes:
    wb = Workbook()

    # Hoja 1: respuestas
    ws = wb.active
    ws.title = "Respuestas"
    ws.append(["Pregunta", "Tipo", "Persona", "Respuesta"])
    for q in fg.questions:
        tipo = "uno_a_uno" if q.tipo == "uno_a_uno" else "general"
        for r in sorted(q.responses, key=lambda x: x.orden):
            ws.append([q.texto, tipo, r.persona_nombre, r.texto])

    # Hoja 2: resumen (último informe si existe)
    ws2 = wb.create_sheet("Resumen")
    if fg.reports:
        last = sorted(fg.reports, key=lambda x: x.generated_at)[-1]
        ws2.append(["Informe (Markdown)"])
        for line in last.contenido_markdown.splitlines():
            ws2.append([line])
    else:
        ws2.append(["Sin informe generado todavía."])

    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()
